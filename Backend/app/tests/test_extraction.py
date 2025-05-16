import os
from datetime import datetime
from docx import Document
from tools.file_reader import extract_text

def create_test_report():
    # Initialize Word document
    doc = Document()
    doc.add_heading('File Extraction Test Report', 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Test cases
    test_files = [
        "SAMPLE LEGAL AGREEMENT.txt",
        "SAMPLE LEGAL AGREEMENT.pdf", 
        "SAMPLE LEGAL AGREEMENT.docx",
        "Book1.xlsx"
    ]
    
    # Create results table
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'File Name'
    hdr_cells[1].text = 'File Type'
    hdr_cells[2].text = 'Status'
    hdr_cells[3].text = 'Details'
    
    # Run tests
    test_dir = os.path.join(os.path.dirname(__file__), "test_files")
    for filename in test_files:
        filepath = os.path.join(test_dir, filename)
        file_ext = os.path.splitext(filename)[1].upper()
        
        row_cells = table.add_row().cells
        row_cells[0].text = filename
        
        try:
            with open(filepath, "rb") as f:
                content = f.read()
                result = extract_text(content, file_ext.lower())
                row_cells[1].text = file_ext
                row_cells[2].text = "PASS"
                row_cells[3].text = f"Extracted {len(result)} characters"
        except Exception as e:
            row_cells[1].text = file_ext
            row_cells[2].text = "FAIL"
            row_cells[3].text = str(e)
    
    # Save report
    report_path = os.path.join(test_dir, "Test_Report.docx")
    doc.save(report_path)
    print(f"Report generated: {report_path}")

if __name__ == "__main__":
    create_test_report()