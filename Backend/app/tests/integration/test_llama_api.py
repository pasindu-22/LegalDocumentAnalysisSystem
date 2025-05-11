import os
import sys
import json
import pytest
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches
from dotenv import load_dotenv

# Configure paths
project_root = Path(__file__).parent.parent.parent.parent
sys.path.append(str(project_root))

# Load environment variables
env_path = project_root / ".env"
if not env_path.exists():
    pytest.skip("Skipping integration tests - .env file not found", allow_module_level=True)
load_dotenv(env_path)

from app.services.llama_analyzer import analyze_legal_text_with_llama

TEST_DOCUMENT = """
NON-DISCLOSURE AGREEMENT
Between Company X ("Disclosing Party") and Consultant Y ("Receiving Party") effective 2024-01-01.

1. CONFIDENTIALITY
1.1 Confidential Information shall mean all non-public business, technical, or financial disclosures.
1.2 Receiving Party shall maintain confidentiality for 3 years post-termination.

2. TERM
2.1 This Agreement shall remain in effect for 2 years from Effective Date.
2.2 Survival: Sections 1 (Confidentiality) shall survive termination.

3. GENERAL PROVISIONS
3.1 Governing Law: California law shall govern this Agreement.
3.2 Entire Agreement: This constitutes the entire understanding between parties.
"""

# --- Test Cases ---
def test_api_response_structure():
    """Test the basic response structure"""
    result = analyze_legal_text_with_llama(TEST_DOCUMENT)
    
    required_fields = {
        "summary": str,
        "key_clauses": list,
        "potential_issues": list,
        "parties_and_obligations": list,
        "processing_time": str
    }
    
    for field, field_type in required_fields.items():
        assert field in result, f"Missing field: {field}"
        assert isinstance(result[field], field_type), f"Invalid type for {field}"

def test_legal_content_extraction():
    """Test extraction of legal provisions"""
    result = analyze_legal_text_with_llama(TEST_DOCUMENT)
    
    # Case-insensitive search in all text fields
    all_text = " ".join([
        result["summary"].lower(),
        " ".join(result["key_clauses"]).lower(),
        " ".join(result["parties_and_obligations"]).lower()
    ])
    
    required_concepts = [
        "confidential",
        "term|duration|effective date",
        "company x|disclosing party",
        "consultant y|receiving party",
        "california"
    ]
    
    for concept in required_concepts:
        assert any(keyword in all_text for keyword in concept.split("|")), \
               f"Concept not found: {concept}"

def test_error_handling():
    """Test error scenarios with flexible validation"""
    # Empty document
    empty_result = analyze_legal_text_with_llama("")
    assert len(empty_result["key_clauses"]) == 0
    assert any(msg in empty_result["summary"].lower() 
              for msg in ["provide text", "empty", "no content"])
    
    # Non-legal text
    nonsense_result = analyze_legal_text_with_llama("Today's weather is sunny")
    assert len(nonsense_result["key_clauses"]) <= 2

@pytest.mark.skipif(
    not os.getenv("GROQ_API_KEY"),
    reason="GROQ_API_KEY not found in environment"
)
def test_api_connectivity():
    """Test API response time"""
    result = analyze_legal_text_with_llama(TEST_DOCUMENT)
    assert float(result["processing_time"][:-1]) < 10.0  # Under 10 seconds

# --- Reporting Functions ---
def run_test_case(name, test_func):
    """Execute and capture test results"""
    try:
        test_func()
        return {
            'name': name,
            'passed': True,
            'message': 'Test passed successfully'
        }
    except AssertionError as e:
        return {
            'name': name,
            'passed': False,
            'message': f"Assertion failed: {str(e)}"
        }
    except Exception as e:
        return {
            'name': name,
            'passed': False,
            'message': f"Unexpected error: {str(e)}"
 # [Previous imports and test cases remain exactly the same...]

# --- Reporting Functions ---
def generate_test_report(test_results):
    """Generate or update test report without error handling changes"""
    report_dir = Path(__file__).parent / "test_files"
    report_dir.mkdir(exist_ok=True)
    report_path = report_dir / "Test_Report.docx"
    
    # Create new document (overwrite existing)
    doc = Document()
    
    # Add report header
    doc.add_heading('Legal Analysis Test Report', 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Add results table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'LightShading-Accent1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Test Case'
    hdr_cells[1].text = 'Status' 
    hdr_cells[2].text = 'Details'
    
    # Populate table
    for test in test_results['details']:
        row_cells = table.add_row().cells
        row_cells[0].text = test['name']
        row_cells[1].text = 'PASSED' if test['passed'] else 'FAILED'
        row_cells[2].text = test['message'][:200]  # Truncate long messages
    
    # Save document
    doc.save(report_path)
    print(f"\nReport saved to: {report_path}\n")

# --- Minimal Execution Block ---
if __name__ == "__main__":
    test_results = {
        'passed': True,
        'details': []
    }
    
    # Manually run tests
    test_results['details'].append(run_test_case("API Structure", test_api_response_structure))
    test_results['details'].append(run_test_case("Content Extraction", test_legal_content_extraction))
    test_results['details'].append(run_test_case("Error Handling", test_error_handling))
    test_results['details'].append(run_test_case("API Connectivity", test_api_connectivity))
    
    # Update overall status
    test_results['passed'] = all(t['passed'] for t in test_results['details'])
    
    # Generate report
    generate_test_report(test_results)
    
    # Exit with status code
    sys.exit(0 if test_results['passed'] else 1)